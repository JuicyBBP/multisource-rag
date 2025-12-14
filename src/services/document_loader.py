"""
Document loader module for MultiSource RAG System.
Handles loading and extracting text from various document formats (PDF, DOCX, TXT, etc.).
"""

from pathlib import Path
from typing import List, Dict, Any
from loguru import logger

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument


class Document:
    """Represents a loaded document with metadata."""

    def __init__(
        self,
        content: str,
        source: str,
        metadata: Dict[str, Any] | None = None,
    ):
        self.content = content
        self.source = source
        self.metadata = metadata or {}

    def __repr__(self) -> str:
        preview = self.content[:100].replace("\n", " ")
        return f"Document(source='{self.source}', content='{preview}...', metadata={self.metadata})"


class DocumentLoader:
    """Loads documents from various file formats."""

    @staticmethod
    def load_pdf_pypdf2(file_path: Path) -> str:
        """Load PDF using PyPDF2.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content

        Raises:
            Exception: If PDF cannot be read
        """
        try:
            text_content = []
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)

            return "\n\n".join(text_content)

        except Exception as e:
            logger.error(f"Error loading PDF with PyPDF2 from {file_path}: {e}")
            raise

    @staticmethod
    def load_pdf_pdfplumber(file_path: Path) -> str:
        """Load PDF using pdfplumber (better for complex PDFs).

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content

        Raises:
            Exception: If PDF cannot be read
        """
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(text)

            return "\n\n".join(text_content)

        except Exception as e:
            logger.error(f"Error loading PDF with pdfplumber from {file_path}: {e}")
            raise

    @staticmethod
    def load_pdf(file_path: Path, use_pdfplumber: bool = True) -> Document:
        """Load PDF file with automatic fallback.

        Args:
            file_path: Path to PDF file
            use_pdfplumber: If True, try pdfplumber first, else PyPDF2 first

        Returns:
            Document object with extracted content

        Raises:
            Exception: If both methods fail
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        if file_path.suffix.lower() != ".pdf":
            raise ValueError(f"Not a PDF file: {file_path}")

        # Try primary method
        primary_method = (
            DocumentLoader.load_pdf_pdfplumber
            if use_pdfplumber
            else DocumentLoader.load_pdf_pypdf2
        )
        fallback_method = (
            DocumentLoader.load_pdf_pypdf2
            if use_pdfplumber
            else DocumentLoader.load_pdf_pdfplumber
        )

        try:
            logger.info(f"Loading PDF with {'pdfplumber' if use_pdfplumber else 'PyPDF2'}: {file_path}")
            content = primary_method(file_path)
            method_used = "pdfplumber" if use_pdfplumber else "PyPDF2"

        except Exception as e:
            logger.warning(
                f"Primary method failed, trying fallback: {e}"
            )
            try:
                content = fallback_method(file_path)
                method_used = "PyPDF2" if use_pdfplumber else "pdfplumber"
                logger.info(f"Fallback method succeeded with {method_used}")

            except Exception as fallback_error:
                logger.error(f"Both PDF loading methods failed: {fallback_error}")
                raise Exception(
                    f"Failed to load PDF with both methods. "
                    f"Primary error: {e}. Fallback error: {fallback_error}"
                )

        # Create metadata
        file_stat = file_path.stat()
        metadata = {
            "file_name": file_path.name,
            "file_path": str(file_path.absolute()),
            "file_size_bytes": file_stat.st_size,
            "loader_method": method_used,
            "num_characters": len(content),
        }

        logger.info(f"Successfully loaded PDF: {file_path.name} ({len(content)} characters)")

        return Document(
            content=content,
            source=str(file_path.absolute()),
            metadata=metadata,
        )

    @staticmethod
    def load_docx(file_path: Path) -> Document:
        """Load DOCX (Word) file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Document object with extracted content

        Raises:
            Exception: If file cannot be read
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        if file_path.suffix.lower() not in [".docx", ".doc"]:
            raise ValueError(f"Not a DOCX file: {file_path}")

        try:
            logger.info(f"Loading DOCX: {file_path}")
            doc = DocxDocument(file_path)

            # Extract text from all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)

            # Create metadata
            file_stat = file_path.stat()
            metadata = {
                "file_name": file_path.name,
                "file_path": str(file_path.absolute()),
                "file_size_bytes": file_stat.st_size,
                "loader_method": "python-docx",
                "num_paragraphs": len(paragraphs),
                "num_characters": len(content),
            }

            logger.info(f"Successfully loaded DOCX: {file_path.name} ({len(content)} characters)")

            return Document(
                content=content,
                source=str(file_path.absolute()),
                metadata=metadata,
            )

        except Exception as e:
            logger.error(f"Error loading DOCX from {file_path}: {e}")
            raise

    @staticmethod
    def load_txt(file_path: Path, encoding: str = "utf-8") -> Document:
        """Load plain text file.

        Args:
            file_path: Path to text file
            encoding: File encoding (default: utf-8)

        Returns:
            Document object with extracted content

        Raises:
            Exception: If file cannot be read
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Text file not found: {file_path}")

        try:
            logger.info(f"Loading TXT: {file_path}")
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()

            # Create metadata
            file_stat = file_path.stat()
            metadata = {
                "file_name": file_path.name,
                "file_path": str(file_path.absolute()),
                "file_size_bytes": file_stat.st_size,
                "loader_method": "text",
                "encoding": encoding,
                "num_characters": len(content),
            }

            logger.info(f"Successfully loaded TXT: {file_path.name} ({len(content)} characters)")

            return Document(
                content=content,
                source=str(file_path.absolute()),
                metadata=metadata,
            )

        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decoding failed, trying latin-1 encoding")
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()

                metadata = {
                    "file_name": file_path.name,
                    "file_path": str(file_path.absolute()),
                    "file_size_bytes": file_path.stat().st_size,
                    "loader_method": "text",
                    "encoding": "latin-1",
                    "num_characters": len(content),
                }

                return Document(
                    content=content,
                    source=str(file_path.absolute()),
                    metadata=metadata,
                )

            except Exception as e:
                logger.error(f"Error loading TXT from {file_path}: {e}")
                raise

        except Exception as e:
            logger.error(f"Error loading TXT from {file_path}: {e}")
            raise

    @staticmethod
    def load(file_path: Path | str) -> Document:
        """Auto-detect file type and load document.

        Args:
            file_path: Path to file

        Returns:
            Document object with extracted content

        Raises:
            ValueError: If file type is not supported
            Exception: If file cannot be read
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return DocumentLoader.load_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return DocumentLoader.load_docx(file_path)
        elif suffix in [".txt", ".md", ".markdown", ".rst"]:
            return DocumentLoader.load_txt(file_path)
        else:
            raise ValueError(
                f"Unsupported file type: {suffix}. "
                f"Supported types: .pdf, .docx, .doc, .txt, .md, .markdown, .rst"
            )

    @staticmethod
    def load_multiple(file_paths: List[Path | str]) -> List[Document]:
        """Load multiple documents.

        Args:
            file_paths: List of file paths

        Returns:
            List of Document objects

        Raises:
            Exception: If any file cannot be read
        """
        documents = []
        for file_path in file_paths:
            try:
                doc = DocumentLoader.load(file_path)
                documents.append(doc)
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
                raise

        logger.info(f"Successfully loaded {len(documents)} documents")
        return documents


if __name__ == "__main__":
    """Test document loader."""
    import sys

    # Configure logger for testing
    logger.remove()
    logger.add(sys.stdout, level="INFO")

    print("=== Document Loader Test ===\n")

    # Test with a simple text file
    test_dir = Path("data/test_docs")
    test_dir.mkdir(parents=True, exist_ok=True)

    # Create a test text file
    test_txt = test_dir / "test.txt"
    test_txt.write_text("This is a test document.\n\nIt has multiple paragraphs.")

    # Load the test file
    doc = DocumentLoader.load(test_txt)
    print(f"Loaded document: {doc}")
    print(f"\nContent preview: {doc.content[:100]}")
    print(f"\nMetadata: {doc.metadata}")

    print("\n✅ Document loader test completed!")
